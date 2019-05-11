'''
为了将word版本的需求文档和测试文档等读入数据库，先将其转换为md。以后，逐步建议使用md来写文档。
本程序只能转换docx文件。
usage：python word2md aaa.docx [abc.md]
'''
# https://www.cnblogs.com/klb561/p/10214195.html 察看如何读入docx
# https://python-docx.readthedocs.io/en/latest/ help https://www.cnblogs.com/zhanghongfeng/p/7043412.html
# 用写字板编辑文件并保存,导致doc的所有style都为NONE,但在写字板显示的格式却没有问题."坑"
'''
doc包含五类对象
doc.paragraphs    #段落集合 ParagraphFormat objects, Run objects, Font objects
doc.tables        #表格集合 _Cell objects, _Row objects, _Column objects, _Rows objects, _Columns objects
doc.sections      #节  集合
doc.styles        #样式集合 _CharacterStyle objects, _ParagraphStyle objects, _TableStyle objects
doc.inline_shapes #内置图形 等等
'''

'''
要搞清楚para、styles的结构，才能处理。
['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Header', 'Footer', 'Document Map', 'Default', 'List Paragraph', 'Balloon Text', 'Heading 4', 'List Header', '样式 仿宋_GB2312 四号 黑色 居中', 'TOC Heading', 'toc 2', 'toc 1', 'toc 3']
'''
import sys
import os
sys.path.append(os.getcwd()+"\\lib")
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

#命令行参数判断，至少要带一个docx文件，2个以上的文件抛弃。
if len(sys.argv) == 1:
	print('usage: python word2md ***.docx [***.md]，如果没有最后一个参数，就产生一个同名的md格式文件。')
	sys.exit(1)
else:
	if len(sys.argv) == 2:
		wordfile=sys.argv[1]
		mdfile=wordfile.split(".")[0]+'.md'
	else:
		wordfile=sys.argv[1]
		mdfile=sys.argv[2].split(".")[0]+'.md'

#读取文档
doc = Document(wordfile)
styles = doc.styles
paragraphs = doc.paragraphs
tables = doc.tables
shapes = doc.inline_shapes

print("paragraph: ",len(paragraphs))
# for ii in range(len(paragraphs)):
	# # print("style: ", paragraphs[ii].style.font.name, paragraphs[ii].style.font.size, paragraphs[ii].style.name, "text: ", paragraphs[ii].text)
	# print("style: ", paragraphs[ii].style, "text: ", paragraphs[ii].text)
	
for pp in paragraphs:
	print(pp.text)	
# print("styles ",len(styles))
# for ii in range(len(styles)):
	# print(styles[ii].type)

# print("tables:  ",len(tables))
# for tt in tables:
for t in tables:
    for r in t.rows:
        for c in r.cells:
            print c.text

# print("shapes: ",len(shapes))

	
#print([s.name for s in styles if s.type==1])
# paragraph_styles = [ s for s in styles if s.type == 1] #WD_STYLE_TYPE.PARAGRAPH  ] 
# print(paragraph_styles) 
# for style in paragraph_styles:  
    # print("style.name: %s" %(style.name)) 
	
# from docx import Document
# # %cd D:YanZan_python2018word
# Docx = Document()
# Docx.add_heading("这是一个一级标题",level=1)
# Docx.add_paragraph("这是一个副级标题","Title")
# A = Docx.add_paragraph("My name is aaa")
# A.add_run("我学习的很快乐，啊哈哈哈哈哈，非常好 Good!!!")
# Docx.add_heading("这是一个二级标题",level=2)
# A = Docx.add_paragraph("这个是二级标题的内容呀")
# B = A.add_run("二级标题里面的正文 继续添加！！！！！！！")
# B.font.bold = True # 同时我要对这些正文进行加粗~~~~
# B.font.size = (20)
# Docx.add_heading("我爱学习Python以下就是python的logo呀",level=3)
# # Docx.add_picture("1.png")
# Docx.add_table(rows=5, cols=5)
# Docx.save("Python.docx")